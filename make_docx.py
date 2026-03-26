import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

md_path = r"C:\Users\Ben R\.gemini\antigravity\scratch\linkedin-ifier\Captspeare_Documentation.md"
out_path = r"C:\Users\Ben R\.gemini\antigravity\scratch\linkedin-ifier\Captspeare_Documentation.docx"

doc = Document()

# Style the document
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

with open(md_path, encoding='utf-8') as f:
    lines = f.readlines()

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_table_from_lines(doc, table_lines):
    # Parse markdown table
    rows = []
    for line in table_lines:
        line = line.strip()
        if re.match(r'^\|[-| ]+\|$', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        if cells:
            rows.append(cells)
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Light List Accent 1'
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < max_cols:
                table.rows[i].cells[j].text = cell
                if i == 0:
                    for para in table.rows[i].cells[j].paragraphs:
                        for run in para.runs:
                            run.bold = True
    doc.add_paragraph()

in_code = False
code_lines = []
in_table = False
table_lines = []

for line in lines:
    stripped = line.rstrip('\n')

    # Code blocks
    if stripped.startswith('```'):
        if not in_code:
            in_code = True
            code_lines = []
        else:
            in_code = False
            p = doc.add_paragraph('\n'.join(code_lines), style='No Spacing')
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x22, 0x80, 0x50)
            doc.add_paragraph()
        continue

    if in_code:
        code_lines.append(stripped)
        continue

    # Tables
    if stripped.startswith('|'):
        in_table = True
        table_lines.append(stripped)
        continue
    elif in_table:
        add_table_from_lines(doc, table_lines)
        in_table = False
        table_lines = []

    # Headings
    if stripped.startswith('### '):
        add_heading(doc, stripped[4:], 3)
    elif stripped.startswith('## '):
        add_heading(doc, stripped[3:], 2)
    elif stripped.startswith('# '):
        add_heading(doc, stripped[2:], 1)
    elif stripped.startswith('---'):
        doc.add_paragraph('─' * 60)
    elif stripped.startswith('- '):
        doc.add_paragraph(stripped[2:], style='List Bullet')
    elif stripped.startswith('  - '):
        doc.add_paragraph(stripped[4:], style='List Bullet 2')
    elif stripped == '':
        doc.add_paragraph()
    else:
        # Bold inline
        p = doc.add_paragraph()
        parts = re.split(r'\*\*(.+?)\*\*', stripped)
        for idx, part in enumerate(parts):
            run = p.add_run(part)
            if idx % 2 == 1:
                run.bold = True

if in_table:
    add_table_from_lines(doc, table_lines)

doc.save(out_path)
print(f"Saved: {out_path}")
